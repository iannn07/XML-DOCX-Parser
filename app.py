import io
import re
import base64
import logging
import os
import json  # Added for image handling
from typing import Dict, List, Optional, BinaryIO
from dataclasses import dataclass
from enum import Enum
from flask import Flask, request, jsonify, abort, send_file
from docx import Document
from docx.enum.text import WD_COLOR
from lxml import etree

# Import image handler - add this line to your imports
from docx_image_handler import DOCXImageHandler

# Configure logging
log_level = os.environ.get("LOG_LEVEL", "INFO")
logging.basicConfig(
    level=getattr(logging, log_level),
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


class FormatType(Enum):
    """Enum for supported format types"""

    HIGHLIGHT = "highlight"
    BOLD = "bold"


@dataclass
class FormatMarker:
    """Data class for format markers"""

    type: FormatType
    start: str
    end: str

    @property
    def pattern(self) -> str:
        """Get regex pattern for this marker"""
        return f"{re.escape(self.start)}|{re.escape(self.end)}"


class DOCXFormatter:
    """
    Main formatter class for DOCX documents.

    Features:
    - Processes formatting markers ({{BOLD_START}}, {{HIGHLIGHT_START}}, etc.)
    - Applies formatting and removes markers from the final document
    - Preserves document structure including:
      * Separate paragraphs (Enter key)
      * Line breaks within paragraphs (Shift+Enter)
      * Table structure
      * Headers and footers
      * Plain Text Content Controls
    - Works with complex nested formatting
    - Maintains spacing and text flow
    """

    # Define format markers
    MARKERS = [
        FormatMarker(
            type=FormatType.HIGHLIGHT,
            start="{{HIGHLIGHT_START}}",
            end="{{HIGHLIGHT_END}}",
        ),
        FormatMarker(type=FormatType.BOLD, start="{{BOLD_START}}", end="{{BOLD_END}}"),
    ]

    def __init__(self):
        self._compile_patterns()
        logger.info(
            f"DOCXFormatter initialized with markers: {[m.start for m in self.MARKERS]}"
        )

    def _compile_patterns(self) -> None:
        """Pre-compile regex patterns for performance"""
        patterns = [marker.pattern for marker in self.MARKERS]
        self.marker_pattern = re.compile(f"({'|'.join(patterns)})")

        # Create lookup maps for O(1) marker identification
        self.start_markers = {m.start: m.type for m in self.MARKERS}
        self.end_markers = {m.end: m.type for m in self.MARKERS}

        # Debug: Test the pattern
        test_text = "Test {{BOLD_START}}bold{{BOLD_END}} text"
        test_parts = self.marker_pattern.split(test_text)
        logger.debug(f"Pattern test - Input: {test_text}")
        logger.debug(f"Pattern test - Parts: {test_parts}")
        logger.debug(
            f"Pattern test - Expected: ['Test ', '{{BOLD_START}}', 'bold', '{{BOLD_END}}', ' text']"
        )

    def format_document(self, input_stream: BinaryIO) -> BinaryIO:
        """
        Main method to format a DOCX document.

        This formatter:
        - Finds and processes formatting markers ({{BOLD_START}}, {{HIGHLIGHT_END}}, etc.)
        - Applies the requested formatting (bold, highlight)
        - Removes the markers from the final text
        - Preserves line breaks (soft returns created with Shift+Enter)
        - Works with both regular text and Plain Text Content Controls

        Args:
            input_stream: Binary stream of input DOCX file

        Returns:
            Binary stream of formatted DOCX file
        """
        try:
            document = Document(input_stream)
        except Exception as e:
            logger.error(f"Error opening document: {e}")
            raise ValueError("Invalid DOCX file format") from e

        # Debug: Print document structure info
        if logger.level == logging.DEBUG:
            self._debug_document_structure(document)

        # Process all paragraphs in the document
        processed_count = self._process_all_paragraphs(document)
        logger.info(f"Processed {processed_count} paragraphs with formatting markers")

        # Save and return the formatted document
        output_stream = io.BytesIO()
        document.save(output_stream)
        output_stream.seek(0)
        return output_stream

    def _debug_document_structure(self, document: Document) -> None:
        """Print debug information about document structure"""
        logger.debug("=== Document Structure Debug ===")
        logger.debug(f"Total paragraphs: {len(document.paragraphs)}")
        logger.debug(f"Total tables: {len(document.tables)}")

        # Sample first few paragraphs
        for i, para in enumerate(document.paragraphs[:5]):
            text = para.text[:100] + "..." if len(para.text) > 100 else para.text
            logger.debug(f"Paragraph {i}: {repr(text)}")

        # Check for content controls in XML
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        sdt_count = len(list(document._element.iter(f"{w_ns}sdt")))
        logger.debug(f"Content controls (SDT) found: {sdt_count}")
        logger.debug("================================")

    def _process_all_paragraphs(self, document: Document) -> int:
        """Process all paragraphs including those in content controls"""
        processed_count = 0

        # Process all paragraphs in the document
        for paragraph in document.paragraphs:
            if self._process_paragraph(paragraph):
                processed_count += 1

        # Process paragraphs in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self._process_paragraph(paragraph):
                            processed_count += 1

        # Process paragraphs in headers and footers
        for section in document.sections:
            # Header
            if hasattr(section, "header") and section.header:
                for paragraph in section.header.paragraphs:
                    if self._process_paragraph(paragraph):
                        processed_count += 1
            # Footer
            if hasattr(section, "footer") and section.footer:
                for paragraph in section.footer.paragraphs:
                    if self._process_paragraph(paragraph):
                        processed_count += 1

        # Special handling for content controls that might not be captured above
        cc_processed = self._process_content_controls_special(document)
        processed_count += cc_processed

        return processed_count

    def _process_content_controls(self, document: Document) -> None:
        """Process paragraphs within plain text content controls"""
        # Define Word namespaces
        WORD_NAMESPACE = (
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        )
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        # Find all content control elements using full namespace
        for sdt in document.element.xpath(f".//{WORD_NAMESPACE}sdt"):
            # Find paragraphs within content controls
            for para_elem in sdt.xpath(f".//{WORD_NAMESPACE}p"):
                # Process this paragraph element
                self._process_paragraph_element(para_elem, WORD_NAMESPACE)

    def _process_paragraph(self, paragraph) -> bool:
        """Process a single paragraph object. Returns True if processed."""
        try:
            full_text = self._get_paragraph_text(paragraph)

            # Skip if no markers found
            if not self._has_markers(full_text):
                return False

            logger.debug(
                f"Processing paragraph with markers: {repr(full_text[:50])}..."
            )

            # Check if paragraph has line breaks
            has_breaks = False
            for run in paragraph.runs:
                run_elem = run._element
                w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                br_elems = run_elem.findall(f".//{w_ns}br")
                if br_elems:
                    has_breaks = True
                    logger.debug(
                        f"Paragraph contains {len(br_elems)} line breaks in run"
                    )

            if has_breaks:
                logger.debug("Using line-break-aware processing")

            # Clear and rebuild runs
            self._rebuild_paragraph_runs(paragraph, full_text)
            return True
        except Exception as e:
            logger.error(f"Error processing paragraph: {e}", exc_info=True)
            return False

    def _process_paragraph_element(self, para_elem, namespace: str) -> None:
        """Process a paragraph element (for content controls)"""
        # Extract text from all runs in the paragraph
        text_parts = []
        for run_elem in para_elem.xpath(f"./{namespace}r"):
            for text_elem in run_elem.xpath(f".//{namespace}t"):
                if text_elem.text:
                    text_parts.append(text_elem.text)

        full_text = "".join(text_parts)

        # Skip if no markers found
        if not self._has_markers(full_text):
            return

        # Clear existing runs
        for run_elem in list(para_elem.xpath(f"./{namespace}r")):
            para_elem.remove(run_elem)

        # Rebuild runs with formatting
        self._rebuild_element_runs(para_elem, full_text)

    def _get_paragraph_text(self, paragraph) -> str:
        """Extract all text from a paragraph"""
        return "".join(run.text or "" for run in paragraph.runs)

    def _has_markers(self, text: str) -> bool:
        """Check if text contains any format markers"""
        has_any = any(
            marker.start in text or marker.end in text for marker in self.MARKERS
        )
        if has_any:
            logger.debug(f"Text has markers: {text[:100]}...")
        return has_any

    def _rebuild_paragraph_runs(self, paragraph, full_text: str) -> None:
        """
        Clear and rebuild paragraph runs with formatting, preserving line breaks.

        This method:
        1. Extracts the content sequence (text and line breaks) from the paragraph
        2. Removes all existing runs
        3. Processes text segments between line breaks, applying formatting based on markers
        4. Re-inserts line breaks at their original positions

        This ensures that soft line breaks (Shift+Enter in Word) are preserved
        while formatting markers are processed and removed.
        """
        para_elem = paragraph._element
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        # Build an ordered sequence of content
        content_sequence = []
        for run in paragraph.runs:
            run_elem = run._element
            # Get all elements in order
            for child in run_elem.iter():
                if child.tag == f"{w_ns}t" and child.text:
                    content_sequence.append(("text", child.text))
                elif child.tag == f"{w_ns}br":
                    content_sequence.append(("break", None))

        logger.debug(
            f"Content sequence: {[(t, c[:20] if c else None) for t, c in content_sequence[:10]]}"
        )

        # Remove all run elements
        for run_elem in list(para_elem.findall(f"{w_ns}r")):
            para_elem.remove(run_elem)

        # Rebuild from sequence
        format_states = {fmt_type: False for fmt_type in FormatType}
        current_text_buffer = []

        for elem_type, content in content_sequence:
            if elem_type == "text":
                current_text_buffer.append(content)
            elif elem_type == "break":
                # Process buffered text
                if current_text_buffer:
                    combined_text = "".join(current_text_buffer)
                    self._process_text_segment_runs(
                        paragraph, combined_text, format_states
                    )
                    current_text_buffer = []

                # Add line break
                self._add_line_break_to_paragraph(paragraph, format_states)

        # Process remaining text
        if current_text_buffer:
            combined_text = "".join(current_text_buffer)
            self._process_text_segment_runs(paragraph, combined_text, format_states)

    def _process_text_segment_runs(
        self, paragraph, text: str, format_states: Dict[FormatType, bool]
    ) -> None:
        """Process a text segment for regular paragraphs"""
        if not text:
            return

        parts = self.marker_pattern.split(text)

        for part in parts:
            if not part:
                continue

            if part in self.start_markers:
                format_states[self.start_markers[part]] = True
                logger.debug(f"Found start marker: {part}")
                continue
            elif part in self.end_markers:
                format_states[self.end_markers[part]] = False
                logger.debug(f"Found end marker: {part}")
                continue
            else:
                logger.debug(f"Adding text run: '{part[:30]}...' with formatting")
                run = paragraph.add_run(part)
                self._apply_formatting(run, format_states)

    def _add_line_break_to_paragraph(
        self, paragraph, format_states: Dict[FormatType, bool]
    ) -> None:
        """Add a line break to a paragraph"""
        run = paragraph.add_run()
        self._apply_formatting(run, format_states)
        run.add_break()  # This adds a <w:br/> element

    def _rebuild_element_runs(self, para_elem, full_text: str) -> None:
        """Rebuild runs for a paragraph element (used for content controls)"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        parts = self.marker_pattern.split(full_text)
        format_states = {fmt_type: False for fmt_type in FormatType}

        for part in parts:
            if not part:
                continue

            # Check if part is a marker
            if part in self.start_markers:
                format_states[self.start_markers[part]] = True
            elif part in self.end_markers:
                format_states[self.end_markers[part]] = False
            else:
                # Create new run element
                run_elem = OxmlElement("w:r")
                run_props = OxmlElement("w:rPr")

                # Apply formatting
                if format_states[FormatType.BOLD]:
                    bold_elem = OxmlElement("w:b")
                    run_props.append(bold_elem)

                if format_states[FormatType.HIGHLIGHT]:
                    highlight_elem = OxmlElement("w:highlight")
                    highlight_elem.set(qn("w:val"), "yellow")
                    run_props.append(highlight_elem)

                if len(run_props):
                    run_elem.append(run_props)

                # Add text
                text_elem = OxmlElement("w:t")
                text_elem.text = part
                # Preserve spaces if needed
                if part.startswith(" ") or part.endswith(" "):
                    text_elem.set(qn("xml:space"), "preserve")
                run_elem.append(text_elem)

                para_elem.append(run_elem)

    def _process_content_controls_special(self, document: Document) -> int:
        """
        Special handling for content controls using XML manipulation.

        Content controls can contain multiple paragraphs. This method ensures
        all paragraphs within each content control are processed while preserving
        their structure and line breaks.
        """
        # Access the document's XML
        doc_xml = document._element

        # Define namespace
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        # Counter for content controls found
        cc_count = 0
        processed_paragraphs = 0

        # Find all structured document tags (content controls)
        for sdt in doc_xml.iter(f"{w_ns}sdt"):
            cc_count += 1
            logger.debug(f"Found content control #{cc_count}")

            # Find the content element
            sdt_content = sdt.find(f"{w_ns}sdtContent")
            if sdt_content is None:
                logger.debug(f"Content control #{cc_count} has no sdtContent element")
                continue

            # Process all paragraphs within the content control
            para_count = 0
            paragraphs_in_cc = list(sdt_content.iter(f"{w_ns}p"))
            logger.debug(
                f"Content control #{cc_count} contains {len(paragraphs_in_cc)} paragraphs"
            )

            for para_elem in paragraphs_in_cc:
                para_count += 1
                logger.debug(
                    f"Processing paragraph {para_count} in content control #{cc_count}"
                )
                if self._process_xml_paragraph(para_elem, w_ns):
                    processed_paragraphs += 1

            if para_count == 0:
                logger.debug(f"No paragraphs found in content control #{cc_count}")

        if cc_count > 0:
            logger.info(
                f"Found {cc_count} content controls, processed {processed_paragraphs} paragraphs with markers"
            )

        return processed_paragraphs

    def _process_xml_paragraph(self, para_elem, namespace: str) -> bool:
        """
        Process a paragraph element directly from XML, preserving line breaks.

        This method handles paragraphs within content controls, ensuring that:
        - Text formatting markers are processed and removed
        - Line breaks (soft returns) are preserved
        - The paragraph structure remains intact
        """
        # Build an ordered list of all content in the paragraph
        content_sequence = []
        run_elems = list(para_elem.findall(f"{namespace}r"))

        # Log paragraph preview for debugging
        para_text_preview = "".join(
            elem.text or "" for elem in para_elem.iter(f"{namespace}t")
        )[:100]
        logger.debug(f"XML paragraph preview: {repr(para_text_preview)}")

        # Process runs in order
        for run_idx, run_elem in enumerate(run_elems):
            # Get all child elements in document order
            elements_in_run = []
            for elem in run_elem.iter():
                if elem.tag == f"{namespace}t" and elem.text:
                    elements_in_run.append(("text", elem.text))
                elif elem.tag == f"{namespace}br":
                    elements_in_run.append(("break", None))

            if elements_in_run:
                logger.debug(f"Run {run_idx} contains: {elements_in_run}")

            content_sequence.extend(elements_in_run)

        # Build full text for marker detection (excluding breaks)
        full_text = "".join(
            content for elem_type, content in content_sequence if elem_type == "text"
        )

        # Skip if no markers found
        if not self._has_markers(full_text):
            return False

        logger.debug(
            f"Processing XML paragraph with markers: {repr(full_text[:50])}..."
        )
        logger.debug(
            f"Content sequence ({len(content_sequence)} elements): {[(t, c[:20] if c else None) for t, c in content_sequence[:10]]}"
        )

        # Remove all existing runs
        for run_elem in run_elems:
            para_elem.remove(run_elem)

        # Process the content sequence
        self._rebuild_from_sequence(para_elem, content_sequence, namespace)

        return True

    def _rebuild_from_sequence(
        self, para_elem, content_sequence, namespace: str
    ) -> None:
        """Rebuild paragraph from a sequence of text and break elements"""
        format_states = {fmt_type: False for fmt_type in FormatType}
        current_text_buffer = []

        for elem_type, content in content_sequence:
            if elem_type == "text":
                # Add text to buffer
                current_text_buffer.append(content)
            elif elem_type == "break":
                # Process any buffered text first
                if current_text_buffer:
                    combined_text = "".join(current_text_buffer)
                    self._process_text_segment(
                        para_elem, combined_text, format_states, namespace
                    )
                    current_text_buffer = []

                # Add the line break
                self._create_xml_line_break(para_elem, format_states, namespace)

        # Process any remaining buffered text
        if current_text_buffer:
            combined_text = "".join(current_text_buffer)
            self._process_text_segment(
                para_elem, combined_text, format_states, namespace
            )

    def _process_text_segment(
        self,
        para_elem,
        text: str,
        format_states: Dict[FormatType, bool],
        namespace: str,
    ) -> None:
        """Process a text segment, applying formatting based on markers"""
        if not text:
            return

        # Split by markers
        parts = self.marker_pattern.split(text)

        for part in parts:
            if not part:
                continue

            # Check if part is a marker
            if part in self.start_markers:
                format_states[self.start_markers[part]] = True
                logger.debug(f"Found start marker: {part}")
                continue
            elif part in self.end_markers:
                # Before ending formatting, ensure any pending text is processed
                format_states[self.end_markers[part]] = False
                logger.debug(f"Found end marker: {part}")
                continue
            else:
                # Create run with current formatting
                logger.debug(
                    f"Creating XML run: '{part[:30]}...' with bold={format_states.get(FormatType.BOLD)}, highlight={format_states.get(FormatType.HIGHLIGHT)}"
                )
                self._create_xml_run(para_elem, part, format_states, namespace)

    def _create_xml_line_break(
        self, para_elem, format_states: Dict[FormatType, bool], namespace: str
    ) -> None:
        """Create a line break element in XML"""
        from lxml import etree

        # Create run element for the line break
        run_elem = etree.SubElement(para_elem, f"{namespace}r")

        # Add run properties if formatting is active
        if any(format_states.values()):
            rPr = etree.SubElement(run_elem, f"{namespace}rPr")

            if format_states.get(FormatType.BOLD, False):
                etree.SubElement(rPr, f"{namespace}b")

            if format_states.get(FormatType.HIGHLIGHT, False):
                highlight = etree.SubElement(rPr, f"{namespace}highlight")
                highlight.set(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                    "yellow",
                )

        # Add line break element
        etree.SubElement(run_elem, f"{namespace}br")

    def _create_xml_run(
        self,
        para_elem,
        text: str,
        format_states: Dict[FormatType, bool],
        namespace: str,
    ) -> None:
        """Create a formatted run in XML"""
        from lxml import etree

        # Create run element
        run_elem = etree.SubElement(para_elem, f"{namespace}r")

        # Add run properties if needed
        if any(format_states.values()):
            rPr = etree.SubElement(run_elem, f"{namespace}rPr")

            if format_states[FormatType.BOLD]:
                etree.SubElement(rPr, f"{namespace}b")

            if format_states[FormatType.HIGHLIGHT]:
                highlight = etree.SubElement(rPr, f"{namespace}highlight")
                highlight.set(f"{namespace}val", "yellow")

        # Add text
        text_elem = etree.SubElement(run_elem, f"{namespace}t")
        if text.startswith(" ") or text.endswith(" "):
            # Preserve spaces
            text_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text_elem.text = text

    def _apply_formatting(self, run, format_states: Dict[FormatType, bool]) -> None:
        """Apply formatting to a run based on current format states"""
        if format_states.get(FormatType.BOLD, False):
            run.bold = True
            logger.debug(f"Applied bold to run: '{run.text}'")
        if format_states.get(FormatType.HIGHLIGHT, False):
            run.font.highlight_color = WD_COLOR.YELLOW
            logger.debug(f"Applied highlight to run: '{run.text}'")


class DOCXFormatterAPI:
    """Flask API for DOCX formatting service"""

    def __init__(self):
        self.app = Flask(__name__)
        self.formatter = DOCXFormatter()
        self.image_handler = DOCXImageHandler()  # Add image handler
        self._setup_routes()

    def _setup_routes(self):
        """Setup Flask routes"""
        self.app.route("/format", methods=["POST"])(self.format_endpoint)
        self.app.route("/format-download", methods=["POST"])(
            self.format_download_endpoint
        )
        self.app.route("/format-xml", methods=["POST"])(self.format_xml_endpoint)
        self.app.route("/format-with-images", methods=["POST"])(
            self.format_with_images_endpoint
        )  # New endpoint
        self.app.route("/format-with-images-download", methods=["POST"])(
            self.format_with_images_download_endpoint
        )  # NEW download endpoint
        self.app.route("/test-doc", methods=["GET"])(self.create_test_document)
        self.app.route("/test-images", methods=["GET"])(
            self.test_images_endpoint
        )  # New endpoint
        self.app.route("/test-doc-markers", methods=["POST"])(
            self.test_document_markers_endpoint
        )  # NEW debug endpoint
        self.app.route("/test-json-structure", methods=["POST"])(
            self.test_json_structure_endpoint
        )  # NEW JSON test endpoint
        self.app.route("/test-format", methods=["GET"])(self.test_format_endpoint)
        self.app.route("/test-linebreaks", methods=["GET"])(
            self.test_linebreaks_endpoint
        )
        self.app.route("/health", methods=["GET"])(self.health_check)

    def format_endpoint(self):
        """Main formatting endpoint"""
        logger.info("Received formatting request")

        # Validate request
        if not request.data:
            logger.warning("Empty request body")
            return abort(400, "Request body is empty. Please upload a .docx file.")

        content_type = request.headers.get("Content-Type", "")
        if not self._is_valid_content_type(content_type):
            logger.warning(f"Invalid Content-Type: {content_type}")
            return abort(415, f"Unsupported Media Type: {content_type}")

        try:
            # Process document
            input_stream = io.BytesIO(request.data)
            output_stream = self.formatter.format_document(input_stream)

            # Read output for base64 encoding
            output_data = output_stream.read()
            output_stream.seek(0)

            # Return response in requested format
            response_data = {
                "body": {
                    "$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "$content": base64.b64encode(output_data).decode("utf-8"),
                }
            }

            logger.info("Formatting completed successfully")
            return jsonify(response_data), 200

        except ValueError as ve:
            logger.error(f"Formatting error: {ve}")
            return abort(400, str(ve))
        except Exception as e:
            logger.error(f"Unexpected error: {e}", exc_info=True)
            return abort(500, "Internal server error during document processing")

    def format_with_images_endpoint(self):
        """Format document with images from AI response - NEW ENDPOINT"""
        logger.info("Received formatting request with images")

        # Check Content-Type
        content_type = request.headers.get("Content-Type", "")
        if not content_type.startswith("application/json"):
            logger.warning(
                f"Invalid Content-Type for /format-with-images: {content_type}"
            )
            return (
                jsonify(
                    {
                        "status": "error",
                        "message": "Content-Type must be application/json",
                        "hint": "Set header 'Content-Type: application/json' and send JSON body",
                    }
                ),
                415,
            )

        try:
            # Get JSON data from request
            data = request.get_json()

            if not data:
                return (
                    jsonify(
                        {
                            "status": "error",
                            "message": "Request body must be valid JSON",
                            "example": {
                                "body": {
                                    "$content": "base64_encoded_docx_content",
                                    "$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    "images": [
                                        {
                                            "marker": "IMAGE_ID",
                                            "data": "base64_encoded_image",
                                            "format": "png",
                                            "width": 400,
                                            "height": 300,
                                            "description": "Image description",
                                        }
                                    ],
                                }
                            },
                        }
                    ),
                    400,
                )

            # Debug logging
            logger.debug(f"Received data keys: {list(data.keys())}")
            if "body" in data:
                logger.debug(f"Body keys: {list(data['body'].keys())}")
                if "images" in data["body"]:
                    logger.debug(
                        f"Number of images in body: {len(data['body']['images'])}"
                    )

            # Extract document and images
            # Support multiple formats for flexibility
            doc_content = None
            images = []

            # Format 1: Body structure with images inside body (NEW FORMAT)
            if "body" in data and isinstance(data["body"], dict):
                doc_content = data["body"].get("$content")
                images = data["body"].get("images", [])
            elif "document" in data:  # legacy support
                doc_content = data.get("document")
                images = data.get("images", [])

            if not images and "images" in data:  # <- add these two lines
                images = data["images"]  # <-

            if not doc_content:
                return abort(
                    400, "Missing document content in body.$content or document"
                )

            if not doc_content:
                return (
                    jsonify(
                        {
                            "status": "error",
                            "message": "Missing document content",
                            "hint": "Provide document in 'body.$content' field",
                            "example": {
                                "body": {
                                    "$content": "base64_encoded_docx",
                                    "images": [{"marker": "...", "data": "..."}],
                                }
                            },
                        }
                    ),
                    400,
                )

            # Decode document content
            if isinstance(doc_content, str):
                doc_bytes = base64.b64decode(doc_content)
            else:
                doc_bytes = doc_content

            # Log where we found the data
            if "body" in data and "images" in data.get("body", {}):
                logger.info(f"Found images in body.images: {len(images)} images")
            else:
                logger.info(f"Found images at root level: {len(images)} images")

            # First, format the document for text markers
            input_stream = io.BytesIO(doc_bytes)
            formatted_stream = self.formatter.format_document(input_stream)
            formatted_bytes = formatted_stream.read()

            # Then, add images if any
            if images:
                logger.info(f"Processing {len(images)} images")
                # Log the image markers for debugging
                for img in images:
                    logger.info(
                        f"Image marker: {img.get('marker')}, format: {img.get('format')}, size: {img.get('width')}x{img.get('height')}"
                    )

                result_bytes = self.image_handler.add_images_to_docx(
                    formatted_bytes, images
                )
                logger.info("Image processing completed")
            else:
                result_bytes = formatted_bytes

            # Return response
            response_data = {
                "status": "success",
                "body": {
                    "$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "$content": base64.b64encode(result_bytes).decode("utf-8"),
                },
                "images_processed": len(images),
            }

            logger.info(f"Formatting completed with {len(images)} images")
            return jsonify(response_data), 200

        except ValueError as ve:
            logger.error(f"Formatting error: {ve}")
            return jsonify({"status": "error", "message": str(ve)}), 400
        except Exception as e:
            logger.error(f"Unexpected error: {e}", exc_info=True)
            return jsonify({"status": "error", "message": "Internal server error"}), 500

    def format_with_images_download_endpoint(self):
        """
        Download endpoint for format-with-images - returns file directly
        Use this endpoint to save the response as a file instead of getting JSON
        """
        logger.info("Received formatting request with images (download mode)")

        # Check Content-Type
        content_type = request.headers.get("Content-Type", "")
        if not content_type.startswith("application/json"):
            logger.warning(
                f"Invalid Content-Type for /format-with-images-download: {content_type}"
            )
            return abort(415, "Content-Type must be application/json")

        try:
            # Get JSON data from request
            data = request.get_json()

            if not data:
                return abort(400, "Request body must be valid JSON")

            # Debug logging
            logger.debug(f"Received data keys: {list(data.keys())}")
            if "body" in data:
                logger.debug(f"Body keys: {list(data['body'].keys())}")
                if "images" in data["body"]:
                    logger.debug(
                        f"Number of images in body: {len(data['body']['images'])}"
                    )

            # Extract document and images
            # Support multiple formats for flexibility
            doc_content = None
            images = []

            # Format 1: Body structure with images inside body (NEW FORMAT)
            if "body" in data and isinstance(data["body"], dict):
                doc_content = data["body"].get("$content")
                images = data["body"].get("images", [])
            elif "document" in data:  # legacy support
                doc_content = data.get("document")
                images = data.get("images", [])

            if not images and "images" in data:  # <- add these two lines
                images = data["images"]  # <-

            if not doc_content:
                return abort(
                    400, "Missing document content in body.$content or document"
                )

            # Decode document content
            if isinstance(doc_content, str):
                doc_bytes = base64.b64decode(doc_content)
            else:
                doc_bytes = doc_content

            # Log where we found the data
            if "body" in data and "images" in data.get("body", {}):
                logger.info(f"Found images in body.images: {len(images)} images")
            else:
                logger.info(f"Found images at root level: {len(images)} images")

            # First, format the document for text markers
            input_stream = io.BytesIO(doc_bytes)
            formatted_stream = self.formatter.format_document(input_stream)
            formatted_bytes = formatted_stream.read()

            # Then, add images if any
            if images:
                logger.info(f"Processing {len(images)} images")
                # Log the image markers for debugging
                for img in images:
                    logger.info(
                        f"Image marker: {img.get('marker')}, format: {img.get('format')}, size: {img.get('width')}x{img.get('height')}"
                    )

                result_bytes = self.image_handler.add_images_to_docx(
                    formatted_bytes, images
                )
                logger.info("Image processing completed")
            else:
                result_bytes = formatted_bytes

            # Create output stream
            output_stream = io.BytesIO(result_bytes)
            output_stream.seek(0)

            logger.info(
                f"Formatting completed with {len(images)} images (download mode)"
            )

            # Return as downloadable file
            return send_file(
                output_stream,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name="formatted_document_with_images.docx",
            )

        except ValueError as ve:
            logger.error(f"Formatting error: {ve}")
            return abort(400, str(ve))
        except Exception as e:
            logger.error(f"Unexpected error: {e}", exc_info=True)
            return abort(500, "Internal server error during document processing")

    def format_download_endpoint(self):
        """
        Temporary endpoint for Postman testing - returns file directly
        Use "Save Response -> Save to a file" in Postman to download
        """
        logger.info("Received formatting request (download mode)")

        # Handle both binary and form data uploads
        file_data = None

        # Check if it's a form upload (from test page)
        if "file" in request.files:
            file = request.files["file"]
            if file.filename == "":
                return abort(400, "No file selected")
            file_data = file.read()
        # Otherwise, assume binary upload (from Postman)
        elif request.data:
            file_data = request.data
        else:
            return abort(400, "No file data received")

        try:
            # Process document
            input_stream = io.BytesIO(file_data)
            output_stream = self.formatter.format_document(input_stream)

            logger.info("Formatting completed successfully (download mode)")

            # Return as downloadable file
            return send_file(
                output_stream,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name="formatted_document.docx",
            )

        except ValueError as ve:
            logger.error(f"Formatting error: {ve}")
            return abort(400, str(ve))
        except Exception as e:
            logger.error(f"Unexpected error: {e}", exc_info=True)
            return abort(500, "Internal server error during document processing")

    def format_xml_endpoint(self):
        """XML response endpoint for testing"""
        logger.info("Received formatting request (XML mode)")

        # Validate request
        if not request.data:
            return self._xml_error("Empty request body")

        try:
            # Process document
            input_stream = io.BytesIO(request.data)
            output_stream = self.formatter.format_document(input_stream)

            # Read output for base64 encoding
            output_data = output_stream.read()

            # Create XML response
            xml_response = f"""<?xml version="1.0" encoding="UTF-8"?>
<response>
    <status>success</status>
    <message>Document formatted successfully</message>
    <document>
        <contentType>application/vnd.openxmlformats-officedocument.wordprocessingml.document</contentType>
        <encoding>base64</encoding>
        <content>{base64.b64encode(output_data).decode('utf-8')}</content>
    </document>
</response>"""

            logger.info("Formatting completed successfully (XML mode)")
            return xml_response, 200, {"Content-Type": "application/xml"}

        except ValueError as ve:
            return self._xml_error(str(ve))
        except Exception as e:
            logger.error(f"Unexpected error: {e}", exc_info=True)
            return self._xml_error("Internal server error")

    def _xml_error(self, message: str):
        """Return XML error response"""
        xml_response = f"""<?xml version="1.0" encoding="UTF-8"?>
<response>
    <status>error</status>
    <message>{message}</message>
</response>"""
        return xml_response, 400, {"Content-Type": "application/xml"}

    def test_images_endpoint(self):
        """Test endpoint to create a document with image markers - NEW ENDPOINT"""
        doc = Document()

        # Add title
        doc.add_heading("Test Document with Image Markers", 0)

        # Add paragraph with image marker
        doc.add_paragraph("This is a test paragraph with an image below:")
        doc.add_paragraph("{{IMAGE:test_image_1}}")
        doc.add_paragraph("Text after the image marker.")

        # Add another section
        doc.add_heading("Section with Multiple Images", 1)
        doc.add_paragraph("First image: {{IMAGE:chart_1}}")
        doc.add_paragraph("Second image: {{IMAGE:diagram_1}}")

        # Add combined markers
        doc.add_paragraph(
            "{{BOLD_START}}Bold text{{BOLD_END}} followed by {{IMAGE:inline_image}}"
        )

        # Save to stream
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="test_document_with_image_markers.docx",
        )

    def test_document_markers_endpoint(self):
        """Debug endpoint to check what markers are in a document"""
        logger.info("Checking document for image markers")

        # Check Content-Type
        content_type = request.headers.get("Content-Type", "")
        if not content_type.startswith("application/json"):
            return (
                jsonify(
                    {
                        "status": "error",
                        "message": "Content-Type must be application/json",
                    }
                ),
                415,
            )

        try:
            # Get JSON data
            data = request.get_json()
            if not data:
                return (
                    jsonify(
                        {
                            "status": "error",
                            "message": "Request body must be valid JSON",
                        }
                    ),
                    400,
                )

            # Extract document content
            doc_content = None
            if "body" in data and isinstance(data["body"], dict):
                if "$content" in data["body"]:
                    doc_content = data["body"]["$content"]
            elif "document" in data:
                doc_content = data.get("document")

            if not doc_content:
                return (
                    jsonify({"status": "error", "message": "Missing document content"}),
                    400,
                )

            # Decode document
            doc_bytes = base64.b64decode(doc_content)
            doc = Document(io.BytesIO(doc_bytes))

            # Find all image markers
            markers_found = []
            import re

            image_pattern = re.compile(r"\{\{IMAGE:([^}]+)\}\}")

            # Check all paragraphs
            for para in doc.paragraphs:
                text = para.text
                matches = image_pattern.findall(text)
                for match in matches:
                    markers_found.append(
                        {
                            "marker": match,
                            "full_marker": f"{{{{IMAGE:{match}}}}}",
                            "paragraph": (
                                text[:100] + "..." if len(text) > 100 else text
                            ),
                        }
                    )

            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text
                            matches = image_pattern.findall(text)
                            for match in matches:
                                markers_found.append(
                                    {
                                        "marker": match,
                                        "full_marker": f"{{{{IMAGE:{match}}}}}",
                                        "in_table": True,
                                        "paragraph": (
                                            text[:100] + "..."
                                            if len(text) > 100
                                            else text
                                        ),
                                    }
                                )

            return (
                jsonify(
                    {
                        "status": "success",
                        "markers_found": markers_found,
                        "total_markers": len(markers_found),
                        "unique_markers": list(set(m["marker"] for m in markers_found)),
                    }
                ),
                200,
            )

        except Exception as e:
            logger.error(f"Error checking markers: {e}", exc_info=True)
            return jsonify({"status": "error", "message": str(e)}), 500

    def test_json_structure_endpoint(self):
        """Test endpoint to verify JSON structure is being received correctly"""
        logger.info("Testing JSON structure")

        try:
            # Get raw data first
            raw_data = request.get_data(as_text=True)
            logger.info(f"Raw data length: {len(raw_data)}")

            # Try to parse JSON
            data = request.get_json()

            result = {
                "status": "success",
                "data_keys": list(data.keys()) if data else None,
                "has_body": "body" in data if data else False,
            }

            if data and "body" in data:
                body = data["body"]
                result["body_keys"] = (
                    list(body.keys()) if isinstance(body, dict) else "Not a dict"
                )

                if isinstance(body, dict) and "images" in body:
                    images = body["images"]
                    result["images_count"] = (
                        len(images) if isinstance(images, list) else "Not a list"
                    )
                    if isinstance(images, list) and len(images) > 0:
                        result["first_image_keys"] = (
                            list(images[0].keys())
                            if isinstance(images[0], dict)
                            else "Not a dict"
                        )
                        result["image_markers"] = [
                            img.get("marker", "NO MARKER")
                            for img in images
                            if isinstance(img, dict)
                        ]

            # Also check alternative structure
            if data and "images" in data:
                result["root_images_count"] = (
                    len(data["images"])
                    if isinstance(data["images"], list)
                    else "Not a list"
                )

            return jsonify(result), 200

        except Exception as e:
            logger.error(f"Error testing JSON: {e}", exc_info=True)
            return (
                jsonify(
                    {"status": "error", "message": str(e), "type": type(e).__name__}
                ),
                500,
            )

    def create_test_document(self):
        """Create a test DOCX with formatting markers"""
        doc = Document()

        # Add title
        doc.add_heading("Test Document for Formatting", 0)

        # Add regular paragraph with markers
        doc.add_paragraph(
            "Regular text: This is {{BOLD_START}}bold text{{BOLD_END}} and this is {{HIGHLIGHT_START}}highlighted text{{HIGHLIGHT_END}}."
        )

        # Add paragraph with line breaks (Shift+Enter)
        para = doc.add_paragraph("2.1 First item")
        run = para.add_run()
        run.add_break()  # Line break
        run.add_text("2.2 Second item")
        run.add_break()  # Line break
        run.add_text("{{HIGHLIGHT_START}}2.3 Highlighted item")
        run.add_break()  # Line break
        run.add_text("2.4 Another highlighted item{{HIGHLIGHT_END}}")
        run.add_break()  # Line break
        run.add_text("2.5 Regular item again")

        # Add paragraph with both
        doc.add_paragraph(
            "Combined: {{BOLD_START}}{{HIGHLIGHT_START}}This is bold and highlighted{{HIGHLIGHT_END}}{{BOLD_END}}."
        )

        # Add nested formatting
        doc.add_paragraph(
            "Nested: {{HIGHLIGHT_START}}This is highlighted with {{BOLD_START}}bold inside{{BOLD_END}} it{{HIGHLIGHT_END}}."
        )

        # Add a table with markers
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell with {{BOLD_START}}bold{{BOLD_END}}"
        table.cell(0, 1).text = (
            "Cell with {{HIGHLIGHT_START}}highlight{{HIGHLIGHT_END}}"
        )
        table.cell(1, 0).text = "{{BOLD_START}}Full cell bold{{BOLD_END}}"
        table.cell(1, 1).text = (
            "{{HIGHLIGHT_START}}Full cell highlight{{HIGHLIGHT_END}}"
        )

        # Add note about content controls
        doc.add_paragraph()
        doc.add_heading("Testing with Content Controls", 2)
        doc.add_paragraph("For testing with Plain Text Content Controls:")
        doc.add_paragraph("1. Save this document")
        doc.add_paragraph("2. Open in Word")
        doc.add_paragraph("3. Go to Developer tab → Insert Plain Text Content Control")
        doc.add_paragraph(
            "4. Type markers inside the content control: {{BOLD_START}}test{{BOLD_END}}"
        )
        doc.add_paragraph("5. Save and upload to test the formatter")

        doc.add_paragraph()
        doc.add_heading("Understanding Paragraphs vs Line Breaks", 2)
        doc.add_paragraph(
            "• Paragraphs (Enter key): Creates separate paragraph elements. Each shows as a separate block."
        )
        doc.add_paragraph(
            "• Line breaks (Shift+Enter): Creates a soft return within the same paragraph."
        )
        doc.add_paragraph(
            "This formatter preserves both structures. If content appears merged, check which type you used."
        )

        # Save to stream
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="test_document_with_markers.docx",
        )

    def test_format_endpoint(self):
        """Test the formatting logic directly"""
        # Create a simple test document
        doc = Document()
        doc.add_paragraph(
            "Test: {{BOLD_START}}bold{{BOLD_END}} and {{HIGHLIGHT_START}}highlight{{HIGHLIGHT_END}}"
        )

        # Save to stream
        input_stream = io.BytesIO()
        doc.save(input_stream)
        input_stream.seek(0)

        # Format it
        try:
            output_stream = self.formatter.format_document(input_stream)

            # Load the result
            result_doc = Document(output_stream)

            # Check the first paragraph
            if result_doc.paragraphs:
                para = result_doc.paragraphs[0]
                result_text = para.text
                runs_info = []
                for run in para.runs:
                    runs_info.append(
                        {
                            "text": run.text,
                            "bold": run.bold,
                            "highlight": run.font.highlight_color,
                        }
                    )

                return jsonify(
                    {
                        "success": True,
                        "original": "Test: {{BOLD_START}}bold{{BOLD_END}} and {{HIGHLIGHT_START}}highlight{{HIGHLIGHT_END}}",
                        "result_text": result_text,
                        "runs": runs_info,
                        "expected_text": "Test: bold and highlight",
                        "test_passed": result_text == "Test: bold and highlight",
                    }
                )
            else:
                return jsonify({"success": False, "error": "No paragraphs in result"})

        except Exception as e:
            logger.error(f"Test format error: {e}", exc_info=True)
            return jsonify({"success": False, "error": str(e)})

    def test_linebreaks_endpoint(self):
        """Test line break handling specifically"""
        # Create a document with line breaks
        doc = Document()
        para = doc.add_paragraph()

        # Add text with line breaks
        run = para.add_run("2.1 First item")
        run.add_break()
        run = para.add_run("2.2 Second item")
        run.add_break()
        run = para.add_run("{{HIGHLIGHT_START}}2.3 Highlighted item")
        run.add_break()
        run = para.add_run("2.4 Another highlighted item{{HIGHLIGHT_END}}")
        run.add_break()
        run = para.add_run("2.5 Regular item")

        # Save and format
        input_stream = io.BytesIO()
        doc.save(input_stream)
        input_stream.seek(0)

        try:
            # Enable debug logging for this test
            original_level = logger.level
            logger.setLevel(logging.DEBUG)

            output_stream = self.formatter.format_document(input_stream)

            # Restore original log level
            logger.setLevel(original_level)

            # Load result and analyze
            result_doc = Document(output_stream)
            if result_doc.paragraphs:
                para = result_doc.paragraphs[0]

                # Analyze runs in detail
                runs_info = []
                break_count = 0

                for i, run in enumerate(para.runs):
                    run_elem = run._element
                    w_ns = (
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                    )
                    br_elems = run_elem.findall(f".//{w_ns}br")

                    run_info = {
                        "index": i,
                        "text": run.text or "(empty)",
                        "has_break": len(br_elems) > 0,
                        "break_count": len(br_elems),
                        "is_bold": run.bold or False,
                        "is_highlighted": run.font.highlight_color == WD_COLOR.YELLOW,
                    }
                    runs_info.append(run_info)
                    break_count += len(br_elems)

                # Expected structure after formatting
                expected_structure = [
                    "2.1 First item[BREAK]",
                    "2.2 Second item[BREAK]",
                    "2.3 Highlighted item[BREAK]",  # Should be highlighted
                    "2.4 Another highlighted item[BREAK]",  # Should be highlighted
                    "2.5 Regular item",
                ]

                return jsonify(
                    {
                        "success": True,
                        "original_text": "2.1 First item\\n2.2 Second item\\n{{HIGHLIGHT_START}}2.3 Highlighted item\\n2.4 Another highlighted item{{HIGHLIGHT_END}}\\n2.5 Regular item",
                        "result_text": para.text,
                        "line_breaks_in_result": break_count,
                        "expected_line_breaks": 4,
                        "runs_count": len(para.runs),
                        "runs_detail": runs_info,
                        "expected_structure": expected_structure,
                        "test_passed": break_count == 4,
                    }
                )

            return jsonify({"success": False, "error": "No paragraphs in result"})

        except Exception as e:
            logger.error(f"Line break test error: {e}", exc_info=True)
            return jsonify({"success": False, "error": str(e)})

    def health_check(self):
        """Health check endpoint"""
        return (
            jsonify(
                {
                    "status": "healthy",
                    "service": "docx-formatter",
                    "features": [
                        "text-formatting",
                        "image-insertion",
                    ],  # Added image feature
                    "endpoints": [
                        "/format",
                        "/format-download",
                        "/format-with-images",
                        "/format-with-images-download",  # NEW endpoint listed
                    ],
                }
            ),
            200,
        )

    def _is_valid_content_type(self, content_type: str) -> bool:
        """Check if content type is valid for DOCX"""
        valid_types = [
            "application/octet-stream",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "multipart/form-data",  # For browser uploads
            "application/json",  # For JSON requests with images
        ]
        return any(content_type.startswith(vt) for vt in valid_types)

    def run(self, host="0.0.0.0", port=5000, debug=False):
        """Run the Flask application"""
        logger.info(f"Starting DOCX Formatter API on {host}:{port}")

        # For production, use a proper WSGI server
        if debug:
            self.app.run(host=host, port=port, debug=debug)
        else:
            # Use waitress for Windows or gunicorn for Linux/Mac
            try:
                from waitress import serve

                serve(self.app, host=host, port=port)
            except ImportError:
                logger.warning(
                    "Waitress not installed, falling back to Flask dev server"
                )
                self.app.run(host=host, port=port, debug=False)


# Create global instance for waitress-serve
api = DOCXFormatterAPI()
app = api.app  # Expose the Flask app instance for waitress-serve


# Quick self-test
def self_test():
    """Run a quick self-test of the formatter"""
    formatter = DOCXFormatter()
    test_text = "Test {{BOLD_START}}bold{{BOLD_END}} and {{HIGHLIGHT_START}}highlight{{HIGHLIGHT_END}}"
    parts = formatter.marker_pattern.split(test_text)

    if logger.level <= logging.DEBUG:
        print("=== FORMATTER SELF-TEST ===")
        print(f"Test text: {test_text}")
        print(f"Split parts: {parts}")
        print(
            f"Expected: ['Test ', '{{BOLD_START}}', 'bold', '{{BOLD_END}}', ' and ', '{{HIGHLIGHT_START}}', 'highlight', '{{HIGHLIGHT_END}}', '']"
        )
        print("===========================")

    # Verify the pattern works correctly
    expected_parts = [
        "Test ",
        "{{BOLD_START}}",
        "bold",
        "{{BOLD_END}}",
        " and ",
        "{{HIGHLIGHT_START}}",
        "highlight",
        "{{HIGHLIGHT_END}}",
        "",
    ]
    if parts != expected_parts:
        logger.error(f"Self-test failed! Expected {expected_parts}, got {parts}")
        raise RuntimeError("Formatter self-test failed")


# Create and run the API
if __name__ == "__main__":
    # Run self-test first
    self_test()

    # When running directly with python app.py
    import sys

    debug_mode = "--debug" in sys.argv
    if debug_mode:
        logger.setLevel(logging.DEBUG)
        logger.info("Running in DEBUG mode")
    api.run(debug=debug_mode)  # Set based on command line arg
